import openai
from docx import Document


class MeetingMinutesGenerator:
    def __init__(self, api_key):
        openai.api_key = api_key
    
    def transcribe_audio(self, audio_file_path):
        with open(audio_file_path, 'rb') as audio_file:
            transcription = openai.audio.transcriptions.create(model="whisper-1", file=audio_file)
        return transcription['text']
    
    def generate_meeting_minutes(self, transcription):
        abstract_summary = self._abstract_summary_extraction(transcription)
        key_points = self._key_points_extraction(transcription)
        action_items = self._action_item_extraction(transcription)
        sentiment = self._sentiment_analysis(transcription)
        return {
            'abstract_summary': abstract_summary,
            'key_points': key_points,
            'action_items': action_items,
            'sentiment': sentiment
        }
    
    def _abstract_summary_extraction(self, transcription):
        response = openai.ChatCompletion.create(
            model="gpt-4",
            temperature=0,
            messages=[
                {"role": "system", "content": "You are a highly skilled, neuro-biologist AI trained in language comprehension and summarization..."},
                {"role": "user", "content": transcription}
            ]
        )
        return response.choices[0].message['content']
    
    def _key_points_extraction(self, transcription):
        response = openai.ChatCompletion.create(
            model="gpt-4",
            temperature=0,
            messages=[
                {"role": "system", "content": "You are a proficient neuro-biologist AI with a specialty in distilling information into key points..."},
                {"role": "user", "content": transcription}
            ]
        )
        return response.choices[0].message['content']
    
    def _action_item_extraction(self, transcription):
        response = openai.ChatCompletion.create(
            model="gpt-4",
            temperature=0,
            messages=[
                {"role": "system", "content": "You are a neuro-biologist AI expert in analyzing conversations and extracting action items..."},
                {"role": "user", "content": transcription}
            ]
        )
        return response.choices[0].message['content']
    
    def _sentiment_analysis(self, transcription):
        response = openai.ChatCompletion.create(
            model="gpt-4",
            temperature=0,
            messages=[
                {"role": "system", "content": "As a neuro-biologist AI with expertise in language and emotion analysis, your task is to analyze the sentiment..."},
                {"role": "user", "content": transcription}
            ]
        )
        return response.choices[0].message['content']
    
    def save_as_docx(self, minutes, filename):
        doc = Document()
        for key, value in minutes.items():
            heading = ' '.join(word.capitalize() for word in key.split('_'))
            doc.add_heading(heading, level=1)
            doc.add_paragraph(value)
            doc.add_paragraph()
        doc.save(filename)

